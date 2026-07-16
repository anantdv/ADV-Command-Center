from __future__ import annotations

from pathlib import Path
from typing import Any

from app.config import settings
from app.core.exceptions import AppError
from app.schemas.document_intake import OCRResult
from app.utils.ocr_layout_extractor import extract_ocr_layout
from app.utils.pdf_text_extractor import extract_pdf_text, is_good_pdf_text


class OCRService:
    async def extract_text(self, intake_id: str, file_path: str, mime_type: str) -> OCRResult:
        document = await self.extract_document_text(file_path, mime_type)
        text = document.get("full_text") or ""
        if not text.strip():
            raise AppError(
                "OCR is not fully configured on this server. Please install OCR dependencies.",
                503,
                {"code": "ocr_dependency_missing"},
            )
        preview = "\n".join(document.get("lines") or [])[:4000] or " ".join(text.split())[:4000]
        return OCRResult(
            intake_id=intake_id,
            extracted_text_preview=preview,
            full_text_available=bool(text),
            confidence=document.get("confidence"),
            page_count=len(document.get("pages") or []) or None,
            source=document.get("source"),
            full_text=text,
            lines=document.get("lines") or [],
            pages=document.get("pages") or [],
            tables=document.get("tables") or [],
            diagnostics=document.get("diagnostics") or {},
        )

    async def extract_document_text(self, file_path: str, mime_type: str | None = None) -> dict[str, Any]:
        if not settings.enable_ocr:
            raise AppError("OCR is disabled.", 503)
        path = Path(file_path)
        diagnostics: dict[str, Any] = {"selected_source": None, "warnings": []}
        if mime_type == "application/pdf" or path.suffix.lower() == ".pdf":
            pdf_text = extract_pdf_text(str(path))
            diagnostics.update(pdf_text.diagnostics)
            if is_good_pdf_text(pdf_text):
                diagnostics["selected_source"] = "pdf_text"
                lines = [line for page in pdf_text.pages for line in page.lines]
                return {
                    "source": "pdf_text",
                    "full_text": pdf_text.full_text,
                    "lines": lines,
                    "pages": [page.to_dict() for page in pdf_text.pages],
                    "tables": pdf_text.tables or [],
                    "confidence": pdf_text.confidence,
                    "diagnostics": diagnostics | {"first_30_lines": lines[:30]},
                }
            ocr = extract_ocr_layout(str(path))
            diagnostics.update({f"ocr_{key}": value for key, value in ocr.diagnostics.items()})
            diagnostics["selected_source"] = "ocr"
            diagnostics["pdf_text_length"] = len(pdf_text.full_text or "")
            diagnostics["ocr_text_length"] = len(ocr.full_text or "")
            diagnostics["first_30_lines"] = ocr.lines[:30]
            return {
                "source": "ocr",
                "full_text": ocr.full_text,
                "lines": ocr.lines,
                "pages": [{"page_number": 1, "text": ocr.full_text, "lines": ocr.lines, "source": "ocr"}],
                "tables": pdf_text.tables or ocr.tables,
                "confidence": 0.55 if ocr.full_text else 0.0,
                "diagnostics": diagnostics,
            }
        elif (mime_type or "").startswith("image/"):
            ocr = extract_ocr_layout(str(path))
            diagnostics.update(ocr.diagnostics)
            diagnostics["selected_source"] = "ocr"
            diagnostics["first_30_lines"] = ocr.lines[:30]
            return {
                "source": "ocr",
                "full_text": ocr.full_text,
                "lines": ocr.lines,
                "pages": [{"page_number": 1, "text": ocr.full_text, "lines": ocr.lines, "source": "ocr"}],
                "tables": ocr.tables,
                "confidence": 0.55 if ocr.full_text else 0.0,
                "diagnostics": diagnostics,
            }
        elif path.suffix.lower() == ".docx":
            text = self._docx_text(path)
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            diagnostics["selected_source"] = "docx_text"
            diagnostics["first_30_lines"] = lines[:30]
            return {
                "source": "docx_text",
                "full_text": text,
                "lines": lines,
                "pages": [{"page_number": 1, "text": text, "lines": lines, "source": "docx_text"}],
                "tables": [],
                "confidence": 0.75 if text else 0.0,
                "diagnostics": diagnostics,
            }
        else:
            raise AppError("Unsupported document format for OCR.", 415)

    @staticmethod
    def _pdf_text(path: Path) -> tuple[str, int]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = min(len(reader.pages), settings.ocr_max_pages)
            text = "\n".join((reader.pages[i].extract_text() or "") for i in range(pages))
            return text, pages
        except Exception:
            return "", 0

    @staticmethod
    def _ocr_pdf(path: Path) -> str:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            pages = convert_from_path(str(path), first_page=1, last_page=settings.ocr_max_pages)
            return "\n".join(pytesseract.image_to_string(page, lang=settings.ocr_language) for page in pages)
        except Exception:
            return ""

    @staticmethod
    def _ocr_image(path: Path) -> str:
        try:
            from PIL import Image
            import pytesseract
            return pytesseract.image_to_string(Image.open(path), lang=settings.ocr_language)
        except Exception:
            return ""

    @staticmethod
    def _docx_text(path: Path) -> str:
        try:
            from docx import Document
            document = Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return ""
